# mcp_main.py
import io
import os
from html import escape
from fastapi import FastAPI
from fastapi import HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# --- Model and App Initialization ---

# Initialize the FastAPI app
app = FastAPI(
    title="Document Summarization MCP",
    description="A microservice to summarize texts and generate a document.",
    version="1.0.0"
)

generator = None
MODEL_NAME = os.getenv("SUMMARIZATION_MODEL", "sshleifer/distilbart-cnn-12-6")

def get_generator():
    """Load the summarizer only when the first summary request is received."""
    global generator
    if generator is None:
        from transformers import pipeline
        generator = pipeline("summarization", model=MODEL_NAME)
    return generator

# --- Pydantic Models for Request Body ---

class SummarizationRequest(BaseModel):
    documents: dict[str, str] = Field(..., description="A dictionary where keys are filenames and values are the document's text content.")
    doc_type: str = Field("pdf", description="The desired output document type. Either 'pdf' or 'docx'.")

# --- Helper Functions for Document Generation ---

def create_pdf_from_summaries(summaries: dict[str, str]) -> io.BytesIO:
    """Generates a PDF document from a dictionary of summaries."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    for filename, summary in summaries.items():
        story.append(Paragraph(escape(f"Summary for: {filename}"), styles['h2']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(escape(summary).replace("\n", "<br/>"), styles['BodyText']))
        story.append(Spacer(1, 24))

    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_from_summaries(summaries: dict[str, str]) -> io.BytesIO:
    """Generates a DOCX document from a dictionary of summaries."""
    buffer = io.BytesIO()
    doc = Document()

    for filename, summary in summaries.items():
        doc.add_heading(f"Summary for: {filename}", level=2)
        doc.add_paragraph(summary)
        doc.add_paragraph()  # Add a little space

    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Helper Function for Text Generation ---

def generate_summary(text: str) -> str:
    """Generate a summary of the given text using the language model."""
    max_input_length = 500
    truncated_text = text[:max_input_length]
    result = get_generator()(truncated_text, max_length=150, truncation=True)
    return result[0]["summary_text"].strip()

# --- API Endpoint ---

@app.post("/summarize-and-create-document/")
async def summarize_and_create(request: SummarizationRequest):
    """
    Receives document texts, summarizes them, and returns a single
    consolidated PDF or DOCX file.
    """
    summaries = {}
    for filename, content in request.documents.items():
        try:
            summary = generate_summary(content)
        except Exception as exc:
            raise HTTPException(status_code=503, detail=f"Summarization model unavailable: {exc}") from exc
        summaries[filename] = summary

    if request.doc_type.lower() == 'pdf':
        buffer = create_pdf_from_summaries(summaries)
        media_type = "application/pdf"
        filename = "summaries.pdf"
    elif request.doc_type.lower() == 'docx':
        buffer = create_docx_from_summaries(summaries)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "summaries.docx"
    else:
        raise HTTPException(status_code=400, detail="Invalid doc_type. Must be 'pdf' or 'docx'.")

    headers = {
        'Content-Disposition': f'attachment; filename="{filename}"'
    }

    return StreamingResponse(buffer, media_type=media_type, headers=headers)

@app.get("/health")
def health_check():
    return {"status": "ok" if generator is not None else "ready", "model": MODEL_NAME}